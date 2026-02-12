"""
Word/DOCX-generator för Jira-tickets
Skapar Word-dokument med alla fält och bilagor
"""
import os
import re
from typing import Dict, Any, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


class WordGenerator:
    """Genererar Word-dokument från Jira-data"""
    
    # Jira-färger
    COLORS = {
        'primary': RGBColor(0, 82, 204),      # Jira blå
        'secondary': RGBColor(23, 43, 77),    # Mörk text
        'accent': RGBColor(0, 135, 90),       # Grön (Done)
        'text_light': RGBColor(94, 108, 132), # Ljusare text
    }
    
    def __init__(self, output_dir: str = None):
        """Initiera Word-generatorn"""
        self.output_dir = output_dir or 'exports'
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_docx(self, issue_data: Dict[str, Any], 
                      attachment_paths: List[str] = None) -> str:
        """
        Generera Word-dokument för en Jira-issue
        
        Args:
            issue_data: Parsad issue-data från JiraClient
            attachment_paths: Lista med sökvägar till nedladdade bilagor
            
        Returns:
            Sökväg till genererat Word-dokument
        """
        issue_key = issue_data['key']
        filename = f"{issue_key}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Sätt dokumentets grundstil
        self._setup_styles(doc)
        
        # Header med issue-typ och nyckel
        self._add_header(doc, issue_data)
        
        # Titel
        title = doc.add_heading(issue_data['summary'], level=1)
        title.runs[0].font.color.rgb = self.COLORS['secondary']
        
        # Status och prioritet
        self._add_status_row(doc, issue_data)
        
        # Separator
        doc.add_paragraph('─' * 60)
        
        # Detaljer
        self._add_details_section(doc, issue_data)
        
        # Beskrivning
        self._add_description_section(doc, issue_data)
        
        # Custom fields
        self._add_custom_fields_section(doc, issue_data)
        
        # Bilagor och bilder
        self._add_attachments_section(doc, issue_data, attachment_paths)
        
        # Subtasks
        self._add_subtasks_section(doc, issue_data)
        
        # Länkar
        self._add_links_section(doc, issue_data)
        
        # Kommentarer
        self._add_comments_section(doc, issue_data)
        
        # Footer
        self._add_footer(doc, issue_data)
        
        # Spara dokument
        doc.save(filepath)
        
        return filepath
    
    def _setup_styles(self, doc: Document):
        """Konfigurera dokumentstilar"""
        # Normal stil
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _add_header(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till dokumenthuvud"""
        issue_type = issue_data['issue_type']['name']
        issue_key = issue_data['key']
        
        p = doc.add_paragraph()
        run1 = p.add_run(f"{issue_type} | ")
        run1.font.color.rgb = self.COLORS['text_light']
        run1.font.size = Pt(12)
        
        run2 = p.add_run(issue_key)
        run2.font.color.rgb = self.COLORS['primary']
        run2.font.bold = True
        run2.font.size = Pt(14)
    
    def _add_status_row(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till status och prioritet"""
        status = issue_data['status']['name']
        priority = issue_data['priority']['name']
        
        p = doc.add_paragraph()
        p.add_run("Status: ").bold = True
        status_run = p.add_run(status)
        
        # Färg baserat på status
        status_lower = status.lower()
        if 'done' in status_lower or 'closed' in status_lower:
            status_run.font.color.rgb = self.COLORS['accent']
        elif 'progress' in status_lower:
            status_run.font.color.rgb = self.COLORS['primary']
        
        p.add_run("    ")
        p.add_run("Prioritet: ").bold = True
        p.add_run(priority)
    
    def _add_details_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till detaljer-sektion"""
        doc.add_heading('Detaljer', level=2)
        
        details = []
        
        if issue_data.get('assignee'):
            details.append(('Tilldelad', issue_data['assignee']['name']))
        else:
            details.append(('Tilldelad', 'Ej tilldelad'))
        
        if issue_data.get('reporter'):
            details.append(('Rapportör', issue_data['reporter']['name']))
        
        if issue_data.get('fix_versions'):
            versions = ', '.join([v['name'] for v in issue_data['fix_versions']])
            details.append(('Fix Versions', versions))
        
        if issue_data.get('components'):
            details.append(('Komponenter', ', '.join(issue_data['components'])))
        
        if issue_data.get('labels'):
            details.append(('Etiketter', ', '.join(issue_data['labels'])))
        
        if issue_data.get('sprints'):
            sprint_names = [s['name'] for s in issue_data['sprints']]
            details.append(('Sprint', ', '.join(sprint_names)))
        
        if issue_data.get('story_points'):
            details.append(('Story Points', str(issue_data['story_points'])))
        
        if issue_data.get('epic'):
            details.append(('Epic', str(issue_data['epic'])))
        
        if issue_data.get('parent'):
            details.append(('Parent', f"{issue_data['parent']['key']} - {issue_data['parent']['summary']}"))
        
        if issue_data.get('created'):
            details.append(('Skapad', self._format_date(issue_data['created'])))
        
        if issue_data.get('updated'):
            details.append(('Uppdaterad', self._format_date(issue_data['updated'])))
        
        # Skapa tabell
        if details:
            table = doc.add_table(rows=len(details), cols=2)
            table.style = 'Table Grid'
            
            for i, (label, value) in enumerate(details):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[1].text = str(value)
            
            doc.add_paragraph()
    
    def _add_description_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till beskrivning"""
        description = issue_data.get('description', '')
        if not description:
            return
        
        doc.add_heading('Beskrivning', level=2)
        
        # Rensa Jira-markup
        clean_text = self._clean_jira_markup(description)
        
        for para in clean_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    
    def _add_custom_fields_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till custom fields"""
        custom_fields = issue_data.get('custom_fields', {})
        if not custom_fields:
            return
        
        skip_patterns = ['story point', 'sprint', 'epic', 'rank', 'flagged']
        filtered = {}
        
        for name, value in custom_fields.items():
            name_lower = name.lower()
            if any(p in name_lower for p in skip_patterns):
                continue
            if value is None or value == '' or value == []:
                continue
            
            if isinstance(value, list):
                value = ', '.join(str(v) for v in value)
            elif isinstance(value, dict):
                value = str(value)
            
            filtered[name] = value
        
        if not filtered:
            return
        
        doc.add_heading('Ytterligare fält', level=2)
        
        table = doc.add_table(rows=len(filtered), cols=2)
        table.style = 'Table Grid'
        
        for i, (name, value) in enumerate(filtered.items()):
            row = table.rows[i]
            row.cells[0].text = name
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = str(value)[:500]
        
        doc.add_paragraph()
    
    def _add_attachments_section(self, doc: Document, issue_data: Dict[str, Any],
                                 attachment_paths: List[str] = None):
        """Lägg till bilagor med bilder"""
        attachments = issue_data.get('attachments', [])
        if not attachments:
            return
        
        doc.add_heading('Bilagor', level=2)
        
        # Skapa mappning
        path_map = {}
        if attachment_paths:
            for path in attachment_paths:
                filename = os.path.basename(path)
                path_map[filename] = path
        
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
        
        for att in attachments:
            filename = att['filename']
            file_ext = os.path.splitext(filename)[1].lower()
            filepath = path_map.get(filename)
            
            # Info om bilaga
            size_kb = att['size'] / 1024
            size_str = f"{size_kb:.1f} KB" if size_kb < 1024 else f"{size_kb/1024:.1f} MB"
            
            p = doc.add_paragraph()
            p.add_run("📎 ").font.size = Pt(12)
            p.add_run(filename).bold = True
            p.add_run(f" ({size_str}) - {att['author']}")
            
            # Visa bild om möjligt
            if file_ext in image_extensions and filepath and os.path.exists(filepath):
                try:
                    doc.add_picture(filepath, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception as e:
                    p = doc.add_paragraph()
                    p.add_run(f"(Kunde inte visa bild: {e})").italic = True
    
    def _add_subtasks_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till subtasks"""
        subtasks = issue_data.get('subtasks', [])
        if not subtasks:
            return
        
        doc.add_heading('Underuppgifter', level=2)
        
        table = doc.add_table(rows=len(subtasks) + 1, cols=3)
        table.style = 'Table Grid'
        
        # Header
        hdr = table.rows[0]
        hdr.cells[0].text = 'Nyckel'
        hdr.cells[1].text = 'Sammanfattning'
        hdr.cells[2].text = 'Status'
        for cell in hdr.cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, st in enumerate(subtasks, 1):
            row = table.rows[i]
            row.cells[0].text = st['key']
            row.cells[1].text = st['summary'][:80]
            row.cells[2].text = st['status']
        
        doc.add_paragraph()
    
    def _add_links_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till länkade ärenden"""
        links = issue_data.get('links', [])
        if not links:
            return
        
        doc.add_heading('Länkade ärenden', level=2)
        
        for link in links:
            p = doc.add_paragraph()
            p.add_run(f"{link['type']} ").font.color.rgb = self.COLORS['text_light']
            p.add_run(link['key']).font.color.rgb = self.COLORS['primary']
            p.add_run(f" - {link['summary'][:60]}")
    
    def _add_comments_section(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till kommentarer"""
        comments = issue_data.get('comments', [])
        if not comments:
            return
        
        doc.add_heading(f'Kommentarer ({len(comments)})', level=2)
        
        for comment in comments:
            p = doc.add_paragraph()
            p.add_run(comment['author']).bold = True
            p.add_run(f" - {self._format_date(comment['created'])}")
            
            body = self._clean_jira_markup(comment['body'])
            doc.add_paragraph(body[:1000])
            doc.add_paragraph()
    
    def _add_footer(self, doc: Document, issue_data: Dict[str, Any]):
        """Lägg till footer"""
        doc.add_paragraph('─' * 60)
        
        export_time = datetime.now().strftime('%Y-%m-%d %H:%M')
        p = doc.add_paragraph()
        run = p.add_run(f"Exporterad: {export_time} | Issue: {issue_data['key']}")
        run.font.size = Pt(9)
        run.font.color.rgb = self.COLORS['text_light']
    
    def _clean_jira_markup(self, text: str) -> str:
        """Rensa Jira-markup"""
        if not text:
            return ''
        
        patterns = [
            (r'\{code[^}]*\}', ''),
            (r'\{noformat\}', ''),
            (r'\{color[^}]*\}', ''),
            (r'\{panel[^}]*\}', ''),
            (r'h[1-6]\.\s*', ''),
            (r'\[([^\]|]+)\|([^\]]+)\]', r'\1'),
            (r'\[([^\]]+)\]', r'\1'),
            (r'\*([^*]+)\*', r'\1'),
            (r'_([^_]+)_', r'\1'),
        ]
        
        result = text
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result)
        
        return result.strip()
    
    def _format_date(self, date_str: str) -> str:
        """Formatera datum"""
        if not date_str:
            return 'N/A'
        
        try:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime('%Y-%m-%d %H:%M')
        except (ValueError, AttributeError):
            return str(date_str)[:19]
